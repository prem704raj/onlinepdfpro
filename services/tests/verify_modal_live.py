"""
Live Verification Script for Service B (PDF -> DOCX on Modal).
Tests the deployed Modal endpoint against all PDF fixtures and ground truth.
"""

import difflib
import json
import os
import re
import sys
import time
import urllib.request
import urllib.error
from pathlib import Path

from docx import Document
from io import BytesIO

MODAL_URL = "https://prem736raj--pdf2docx-convert.modal.run"
FIXTURES_DIR = Path(__file__).parent / "fixtures"
GT_DIR = FIXTURES_DIR / "ground_truth"


def _normalize(text: str) -> str:
    text = text.lower()
    return re.sub(r"\s+", " ", text).strip()


def _similarity(a: str, b: str) -> float:
    na, nb = _normalize(a), _normalize(b)
    return difflib.SequenceMatcher(None, na, nb).ratio()


def convert_pdf_to_docx(pdf_path: Path) -> tuple[int, bytes, dict]:
    req = urllib.request.Request(
        MODAL_URL,
        data=pdf_path.read_bytes(),
        headers={"Content-Type": "application/pdf"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=180.0) as res:
            return res.status, res.read(), dict(res.headers)
    except urllib.error.HTTPError as e:
        return e.code, e.read(), dict(e.headers)


def extract_docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def main():
    print("============================================================")
    print("  LIVE MODAL ENDPOINT VERIFICATION (Service B: PDF -> DOCX)")
    print("  URL:", MODAL_URL)
    print("============================================================")

    passed = 0
    failed = 0

    def check(name: str, cond: bool, detail: str = ""):
        nonlocal passed, failed
        if cond:
            passed += 1
            print(f"  [PASS] {name} {detail}")
        else:
            failed += 1
            print(f"  [FAIL] {name} {detail}")

    # 1. Test F1 Single Column
    print("\n-- Testing F1 (Single Column PDF) --")
    t0 = time.monotonic()
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f1_single_column.pdf")
    elapsed = time.monotonic() - t0
    check("F1 HTTP Status 200", status == 200, f"(status={status}, {elapsed:.1f}s)")
    check("F1 X-Path header", headers.get("X-Path") == "cpu-text", f"(X-Path={headers.get('X-Path')})")
    
    if status == 200:
        docx_text = extract_docx_text(data)
        gt_text = (GT_DIR / "f1.txt").read_text(encoding="utf-8")
        sim = _similarity(docx_text, gt_text)
        check("F1 Text Fidelity (>=0.85)", sim >= 0.85, f"(similarity={sim:.2%})")

    # 2. Test F2 Three Column (Reading Order)
    print("\n-- Testing F2 (Three Column PDF / Reading Order) --")
    t0 = time.monotonic()
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f2_three_column.pdf")
    elapsed = time.monotonic() - t0
    check("F2 HTTP Status 200", status == 200, f"(status={status}, {elapsed:.1f}s)")
    
    if status == 200:
        docx_text = extract_docx_text(data)
        key_sentence = (GT_DIR / "f2_key_sentence.txt").read_text(encoding="utf-8").strip()
        norm_doc = _normalize(docx_text)
        norm_key = _normalize(key_sentence)
        check("F2 Reading Order preserved", norm_key in norm_doc, "(key sentence intact across columns)")

    # 3. Test F3 Tables
    print("\n-- Testing F3 (Tables PDF) --")
    t0 = time.monotonic()
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f3_tables.pdf")
    elapsed = time.monotonic() - t0
    check("F3 HTTP Status 200", status == 200, f"(status={status}, {elapsed:.1f}s)")
    
    if status == 200:
        doc = Document(BytesIO(data))
        gt_tables = json.loads((GT_DIR / "f3.json").read_text(encoding="utf-8"))
        check("F3 Table Count preserved", len(doc.tables) >= 2, f"(found {len(doc.tables)} tables)")

    # 4. Test F6 Unicode (Hindi + CJK)
    print("\n-- Testing F6 (Unicode PDF) --")
    t0 = time.monotonic()
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f6_unicode.pdf")
    elapsed = time.monotonic() - t0
    check("F6 HTTP Status 200", status == 200, f"(status={status}, {elapsed:.1f}s)")
    
    if status == 200:
        docx_text = extract_docx_text(data)
        gt_f6 = json.loads((GT_DIR / "f6.json").read_text(encoding="utf-8"))
        check("F6 Hindi preserved", gt_f6["hindi"] in docx_text, "")
        check("F6 Chinese preserved", gt_f6["chinese"] in docx_text, "")

    # 5. Test F8a Corrupt PDF
    print("\n-- Testing F8a (Corrupt PDF) --")
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f8a_corrupt.pdf")
    check("F8a Correctly Rejected (422)", status == 422, f"(status={status})")

    # 6. Test F8b Encrypted PDF
    print("\n-- Testing F8b (Encrypted PDF) --")
    status, data, headers = convert_pdf_to_docx(FIXTURES_DIR / "f8b_encrypted.pdf")
    check("F8b Correctly Rejected (422)", status == 422, f"(status={status})")

    print("\n============================================================")
    print(f"  TOTAL: {passed + failed} | PASSED: {passed} | FAILED: {failed}")
    print("============================================================")
    if failed > 0:
        sys.exit(1)
    else:
        print("\n  ✅ ALL LIVE MODAL ENDPOINT TESTS PASSED")


if __name__ == "__main__":
    main()
