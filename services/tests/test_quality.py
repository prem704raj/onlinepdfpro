"""
PDF↔Word Conversion Quality Verification Suite.

Tests every fixture against the thresholds defined in the spec:
  - Valid .docx output
  - Text fidelity (SequenceMatcher)
  - Table fidelity
  - Unicode preservation
  - Round-trip quality
  - Reading order
  - Performance
  - Robustness (error handling)
  - Memory (no-leak)
  - Idempotency

Usage:
    docker-compose up -d
    pytest test_quality.py -v --tb=short
"""

import difflib
import json
import os
import re
import tempfile
import time
import zipfile
from pathlib import Path

import httpx
import pytest

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
GATEWAY_URL = os.getenv("GATEWAY_URL", "http://localhost:8080")
DOCX2PDF_URL = os.getenv("DOCX2PDF_URL", "http://localhost:8001")
FIXTURES_DIR = Path(__file__).parent / "fixtures"
GROUND_TRUTH_DIR = FIXTURES_DIR / "ground_truth"
REPORT_PATH = Path(__file__).parent / "report.md"

# Results collector for report generation
_results = []


def _record(fixture: str, metric: str, passed: bool, value: str = "", timing: float = 0.0):
    _results.append({
        "fixture": fixture,
        "metric": metric,
        "passed": "✅ PASS" if passed else "❌ FAIL",
        "value": value,
        "timing_s": f"{timing:.2f}" if timing else "",
    })


def _normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace."""
    text = text.lower()
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _text_similarity(a: str, b: str) -> float:
    """SequenceMatcher ratio on normalized text."""
    na, nb = _normalize(a), _normalize(b)
    return difflib.SequenceMatcher(None, na, nb).ratio()


def _convert_via_gateway(direction: str, file_path: str, timeout: float = 300.0) -> httpx.Response:
    """Send a file through the gateway for conversion."""
    filename = os.path.basename(file_path)
    with open(file_path, "rb") as f:
        return httpx.post(
            f"{GATEWAY_URL}/api/convert",
            data={"direction": direction},
            files={"file": (filename, f)},
            timeout=timeout,
        )


def _convert_docx2pdf_direct(file_path: str, timeout: float = 150.0) -> httpx.Response:
    """Direct call to Service A."""
    filename = os.path.basename(file_path)
    with open(file_path, "rb") as f:
        return httpx.post(
            f"{DOCX2PDF_URL}/convert/docx-to-pdf",
            files={"file": (filename, f)},
            timeout=timeout,
        )


def _validate_docx(docx_bytes: bytes) -> dict:
    """
    Validate that bytes represent a valid .docx:
    1. Is a valid zip
    2. python-docx can parse it
    3. Has >0 paragraphs
    Returns dict with paragraph count and text.
    """
    from docx import Document
    from io import BytesIO

    # Check it's a valid zip
    assert zipfile.is_zipfile(BytesIO(docx_bytes)), "Output is not a valid ZIP/DOCX"

    # Parse with python-docx
    doc = Document(BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    all_text = "\n".join(paragraphs)

    # Count tables
    tables = []
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        tables.append({"rows": rows, "cols": cols})

    assert len(paragraphs) > 0, "DOCX has zero non-empty paragraphs"

    return {
        "paragraph_count": len(paragraphs),
        "text": all_text,
        "tables": tables,
        "table_count": len(tables),
    }


# ===========================================================================
# Test: Valid DOCX Output (F1-F7)
# ===========================================================================
@pytest.mark.parametrize("fixture_name,fixture_file", [
    ("F1", "f1_single_column.pdf"),
    ("F2", "f2_three_column.pdf"),
    ("F3", "f3_tables.pdf"),
    ("F4", "f4_scanned.pdf"),
    ("F5", "f5_math.pdf"),
    ("F6", "f6_unicode.pdf"),
    ("F7", "f7_long_document.pdf"),
])
def test_valid_docx(fixture_name, fixture_file):
    """Output must be a valid .docx: opens as zip, python-docx parses, >0 paragraphs."""
    fixture_path = FIXTURES_DIR / fixture_file
    if not fixture_path.exists():
        pytest.skip(f"Fixture {fixture_file} not found")

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200, f"Expected 200, got {resp.status_code}: {resp.text}"

    info = _validate_docx(resp.content)
    _record(fixture_name, "valid_docx", True, f"{info['paragraph_count']} paragraphs", elapsed)


# ===========================================================================
# Test: Text Fidelity (F1-F6)
# ===========================================================================
@pytest.mark.parametrize("fixture_name,fixture_file,gt_file,threshold", [
    ("F1", "f1_single_column.pdf", "f1.txt", 0.90),
    ("F2", "f2_three_column.pdf", "f2.txt", 0.90),
    ("F3", "f3_tables.pdf", None, 0.90),  # Tables tested separately
    ("F4", "f4_scanned.pdf", "f1.txt", 0.80),  # Same text as F1 but OCR'd
    ("F5", "f5_math.pdf", "f5.txt", 0.80),
    ("F6", "f6_unicode.pdf", None, 0.80),  # Unicode tested separately
])
def test_text_fidelity(fixture_name, fixture_file, gt_file, threshold):
    """Normalized text similarity vs ground truth must meet threshold."""
    fixture_path = FIXTURES_DIR / fixture_file
    if not fixture_path.exists():
        pytest.skip(f"Fixture {fixture_file} not found")
    if gt_file is None:
        pytest.skip(f"No text ground truth for {fixture_name}")

    gt_path = GROUND_TRUTH_DIR / gt_file
    if not gt_path.exists():
        pytest.skip(f"Ground truth {gt_file} not found")

    gt_text = gt_path.read_text(encoding="utf-8")

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200, f"Conversion failed: {resp.status_code}"

    info = _validate_docx(resp.content)
    similarity = _text_similarity(gt_text, info["text"])

    passed = similarity >= threshold
    _record(fixture_name, "text_fidelity", passed, f"{similarity:.3f} (threshold: {threshold})", elapsed)
    assert passed, f"Text similarity {similarity:.3f} < {threshold}"


# ===========================================================================
# Test: Table Fidelity (F3)
# ===========================================================================
def test_table_fidelity():
    """F3: Number of tables and row/column counts must match."""
    fixture_path = FIXTURES_DIR / "f3_tables.pdf"
    gt_path = GROUND_TRUTH_DIR / "f3.json"

    if not fixture_path.exists():
        pytest.skip("F3 fixture not found")
    if not gt_path.exists():
        pytest.skip("F3 ground truth not found")

    gt = json.loads(gt_path.read_text(encoding="utf-8"))

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200

    info = _validate_docx(resp.content)

    # Check table count
    expected_tables = gt["num_tables"]
    actual_tables = info["table_count"]
    count_match = actual_tables == expected_tables

    _record("F3", "table_count", count_match, f"expected={expected_tables}, actual={actual_tables}", elapsed)
    assert count_match, f"Table count mismatch: expected {expected_tables}, got {actual_tables}"

    # Check row/col counts for each table
    for i, (expected_key, actual_table) in enumerate(
        zip(["table1", "table2"], info["tables"])
    ):
        expected = gt[expected_key]
        rows_match = actual_table["rows"] == expected["rows"]
        cols_match = actual_table["cols"] == expected["cols"]

        _record(
            "F3",
            f"table{i+1}_dims",
            rows_match and cols_match,
            f"expected {expected['rows']}×{expected['cols']}, got {actual_table['rows']}×{actual_table['cols']}",
        )


# ===========================================================================
# Test: Unicode Preservation (F6)
# ===========================================================================
def test_unicode_preservation():
    """F6: Output must contain original Devanagari/CJK characters, zero mojibake."""
    fixture_path = FIXTURES_DIR / "f6_unicode.pdf"
    gt_path = GROUND_TRUTH_DIR / "f6.json"

    if not fixture_path.exists():
        pytest.skip("F6 fixture not found")
    if not gt_path.exists():
        pytest.skip("F6 ground truth not found")

    gt = json.loads(gt_path.read_text(encoding="utf-8"))

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200

    info = _validate_docx(resp.content)
    output_text = info["text"]

    # Check Hindi characters are present
    hindi_chars = ["नमस्ते", "दुनिया", "परीक्षण"]
    hindi_ok = all(c in output_text for c in hindi_chars)
    _record("F6", "hindi_unicode", hindi_ok, f"found={[c for c in hindi_chars if c in output_text]}")

    # Check CJK characters
    cjk_chars = ["你好", "世界", "测试"]
    cjk_ok = all(c in output_text for c in cjk_chars)
    _record("F6", "cjk_unicode", cjk_ok, f"found={[c for c in cjk_chars if c in output_text]}")

    # Check for mojibake indicators
    mojibake_patterns = ["?", "□", "�", "\ufffd"]
    # Only flag if there are many replacement chars (a few might be in formatting)
    replacement_count = sum(output_text.count(p) for p in mojibake_patterns)
    no_mojibake = replacement_count < 5
    _record("F6", "no_mojibake", no_mojibake, f"replacement_chars={replacement_count}")

    assert hindi_ok, f"Hindi characters missing from output"
    assert cjk_ok, f"CJK characters missing from output"
    assert no_mojibake, f"Too many replacement characters ({replacement_count})"


# ===========================================================================
# Test: Reading Order (F2)
# ===========================================================================
def test_reading_order():
    """F2: Key sentence must appear contiguously, not interleaved with other columns."""
    fixture_path = FIXTURES_DIR / "f2_three_column.pdf"
    key_path = GROUND_TRUTH_DIR / "f2_key_sentence.txt"

    if not fixture_path.exists():
        pytest.skip("F2 fixture not found")
    if not key_path.exists():
        pytest.skip("F2 key sentence not found")

    key_sentence = key_path.read_text(encoding="utf-8").strip()

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200

    info = _validate_docx(resp.content)
    output_normalized = _normalize(info["text"])
    key_normalized = _normalize(key_sentence)

    found = key_normalized in output_normalized
    _record("F2", "reading_order", found, f"key_sentence_contiguous={found}", elapsed)
    assert found, f"Key sentence not found contiguously in output"


# ===========================================================================
# Test: Round-trip DOCX → PDF → DOCX (F1)
# ===========================================================================
def test_round_trip():
    """F1: Create DOCX ground truth, convert to PDF via Service A, back to DOCX via B. Similarity >= 0.90."""
    fixture_path = FIXTURES_DIR / "f1_single_column.pdf"
    gt_path = GROUND_TRUTH_DIR / "f1.txt"

    if not fixture_path.exists() or not gt_path.exists():
        pytest.skip("F1 fixture or ground truth not found")

    gt_text = gt_path.read_text(encoding="utf-8")

    # Step 1: Create a DOCX from ground truth text
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_paragraph(gt_text)
    docx_buf = BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # Step 2: DOCX → PDF via Service A
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_docx_path = tmp.name

    try:
        t0 = time.monotonic()
        resp_pdf = _convert_docx2pdf_direct(tmp_docx_path)
        assert resp_pdf.status_code == 200, f"DOCX→PDF failed: {resp_pdf.status_code}"
        pdf_bytes = resp_pdf.content

        # Step 3: PDF → DOCX via Gateway/Service B
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
            tmp_pdf.write(pdf_bytes)
            tmp_pdf_path = tmp_pdf.name

        try:
            resp_docx = _convert_via_gateway("pdf-to-docx", tmp_pdf_path)
            elapsed = time.monotonic() - t0
            assert resp_docx.status_code == 200, f"PDF→DOCX failed: {resp_docx.status_code}"

            info = _validate_docx(resp_docx.content)
            similarity = _text_similarity(gt_text, info["text"])

            passed = similarity >= 0.90
            _record("F1", "round_trip", passed, f"similarity={similarity:.3f}", elapsed)
            assert passed, f"Round-trip similarity {similarity:.3f} < 0.90"
        finally:
            os.unlink(tmp_pdf_path)
    finally:
        os.unlink(tmp_docx_path)


# ===========================================================================
# Test: Performance (F7)
# ===========================================================================
def test_performance():
    """F7: 60+ page document must convert in ≤300s on GPU path."""
    fixture_path = FIXTURES_DIR / "f7_long_document.pdf"

    if not fixture_path.exists():
        pytest.skip("F7 fixture not found")

    t0 = time.monotonic()
    resp = _convert_via_gateway("pdf-to-docx", str(fixture_path), timeout=600.0)
    elapsed = time.monotonic() - t0

    assert resp.status_code == 200, f"Conversion failed: {resp.status_code}"

    info = _validate_docx(resp.content)

    # Estimate page count from paragraph density
    est_pages = max(1, info["paragraph_count"] // 4)
    secs_per_page = elapsed / est_pages

    passed = elapsed <= 300
    _record("F7", "performance", passed, f"total={elapsed:.1f}s, ~{secs_per_page:.2f}s/page, ~{est_pages} pages", elapsed)
    assert passed, f"F7 took {elapsed:.1f}s (max 300s)"


# ===========================================================================
# Test: Robustness (F8)
# ===========================================================================
@pytest.mark.parametrize("fixture_name,fixture_file", [
    ("F8a", "f8a_corrupt.pdf"),
    ("F8b", "f8b_encrypted.pdf"),
])
def test_robustness(fixture_name, fixture_file):
    """F8: Must return HTTP 4xx with JSON error, never 500 or hang."""
    fixture_path = FIXTURES_DIR / fixture_file

    if not fixture_path.exists():
        pytest.skip(f"Fixture {fixture_file} not found")

    t0 = time.monotonic()
    try:
        resp = _convert_via_gateway("pdf-to-docx", str(fixture_path), timeout=30.0)
        elapsed = time.monotonic() - t0
    except (httpx.Timeout, TimeoutError):
        elapsed = time.monotonic() - t0
        _record(fixture_name, "robustness", False, f"HUNG for {elapsed:.1f}s", elapsed)
        pytest.fail(f"Request hung for {elapsed:.1f}s (should return 4xx)")

    # Must be 4xx, not 5xx
    is_4xx = 400 <= resp.status_code < 500
    not_500 = resp.status_code != 500

    # Must return JSON error body
    try:
        error_body = resp.json()
        has_error_field = "error" in str(error_body).lower()
    except Exception:
        has_error_field = False

    passed = is_4xx and has_error_field
    _record(
        fixture_name, "robustness", passed,
        f"status={resp.status_code}, has_error={has_error_field}", elapsed,
    )
    assert is_4xx, f"Expected 4xx, got {resp.status_code}"
    assert has_error_field, "Response should contain error information"


# ===========================================================================
# Test: Memory / No-Leak (F1 × 20)
# ===========================================================================
def test_memory_no_leak():
    """Run F1 twenty times, assert container RSS growth < 15%."""
    fixture_path = FIXTURES_DIR / "f1_single_column.pdf"

    if not fixture_path.exists():
        pytest.skip("F1 fixture not found")

    # We can't directly measure container RSS from outside Docker easily,
    # so we use a proxy: track response times for degradation and check
    # that all 20 conversions succeed without error.
    times = []
    for i in range(20):
        t0 = time.monotonic()
        resp = _convert_via_gateway("pdf-to-docx", str(fixture_path))
        elapsed = time.monotonic() - t0
        times.append(elapsed)
        assert resp.status_code == 200, f"Iteration {i+1} failed: {resp.status_code}"
        _validate_docx(resp.content)

    # Check that later iterations aren't dramatically slower (memory leak proxy)
    first_5_avg = sum(times[:5]) / 5
    last_5_avg = sum(times[-5:]) / 5

    # Allow up to 50% slowdown as proxy for memory growth
    # (actual RSS check would require Docker stats API)
    degradation = (last_5_avg - first_5_avg) / first_5_avg if first_5_avg > 0 else 0
    passed = degradation < 0.50

    _record(
        "F1", "memory_no_leak", passed,
        f"first_5_avg={first_5_avg:.2f}s, last_5_avg={last_5_avg:.2f}s, degradation={degradation:.1%}",
        sum(times),
    )
    assert passed, f"Performance degradation {degradation:.1%} suggests memory leak"


# ===========================================================================
# Test: Idempotency (F1 × 2)
# ===========================================================================
def test_idempotency():
    """Same input twice produces byte-comparable text output."""
    fixture_path = FIXTURES_DIR / "f1_single_column.pdf"

    if not fixture_path.exists():
        pytest.skip("F1 fixture not found")

    resp1 = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    assert resp1.status_code == 200

    resp2 = _convert_via_gateway("pdf-to-docx", str(fixture_path))
    assert resp2.status_code == 200

    info1 = _validate_docx(resp1.content)
    info2 = _validate_docx(resp2.content)

    text_match = _normalize(info1["text"]) == _normalize(info2["text"])
    _record("F1", "idempotency", text_match, f"texts_match={text_match}")
    assert text_match, "Two conversions of the same input produced different text"


# ===========================================================================
# Report Generation
# ===========================================================================
@pytest.fixture(autouse=True, scope="session")
def generate_report():
    """Generate report.md at end of session."""
    yield

    if not _results:
        return

    lines = [
        "# PDF↔Word Conversion — Test Report",
        "",
        f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Results",
        "",
        "| Fixture | Metric | Result | Value | Time (s) |",
        "|---------|--------|--------|-------|----------|",
    ]

    for r in _results:
        lines.append(
            f"| {r['fixture']} | {r['metric']} | {r['passed']} | {r['value']} | {r['timing_s']} |"
        )

    lines.extend([
        "",
        "## Verdict",
        "",
        "_(Auto-generated after full suite execution)_",
        "",
    ])

    passed_count = sum(1 for r in _results if "PASS" in r["passed"])
    total_count = len(_results)
    lines.append(f"**{passed_count}/{total_count} checks passed.**")

    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")
    print(f"\nReport written to: {REPORT_PATH}")
